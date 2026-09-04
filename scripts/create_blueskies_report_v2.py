from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = '/home/vaishnav/airgrid/Blue_Skies_Hack_2026_AirGrid_Final_Report.docx'
GREEN='123C3A'; MINT='EAF4F1'; GRAY='F5F7F6'; BORDER='D5DEDC'

def cell_shade(c, color):
    props=c._tc.get_or_add_tcPr(); el=OxmlElement('w:shd'); el.set(qn('w:fill'),color); props.append(el)
def cell_border(c):
    props=c._tc.get_or_add_tcPr(); borders=OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        el=OxmlElement('w:'+edge); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'5'); el.set(qn('w:color'),BORDER); borders.append(el)
    props.append(borders)
def cell(c,text,bold=False,white=False,size=9.3):
    c.text=''; p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(str(text)); r.bold=bold; r.font.size=Pt(size)
    if white:r.font.color.rgb=RGBColor(255,255,255)
    c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; cell_border(c)
def tbl(doc, heads, rows):
    t=doc.add_table(rows=1,cols=len(heads)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(heads): cell_shade(t.rows[0].cells[i],GREEN); cell(t.rows[0].cells[i],h,True,True)
    for n,row in enumerate(rows):
        for i,x in enumerate(row):
            if n%2: cell_shade(t.add_row().cells[i] if False else t.rows[-1].cells[i],GRAY)
        rr=t.add_row().cells
        for i,x in enumerate(row):
            if n%2: cell_shade(rr[i],MINT)
            cell(rr[i],x)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
def h(doc,text,level=1):
    p=doc.add_heading(text,level); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(5)
    for r in p.runs:r.font.color.rgb=RGBColor(0,0,0)
def p(doc,text,lead=None):
    q=doc.add_paragraph(); q.paragraph_format.space_after=Pt(6); q.paragraph_format.line_spacing=1.12
    if lead:r=q.add_run(lead); r.bold=True
    q.add_run(text)
def bullets(doc,lines):
    for line in lines:
        q=doc.add_paragraph(style='List Bullet'); q.paragraph_format.space_after=Pt(2); q.add_run(line)
def caption(doc,text):
    q=doc.add_paragraph();q.alignment=WD_ALIGN_PARAGRAPH.CENTER;q.paragraph_format.space_after=Pt(8);r=q.add_run(text);r.italic=True;r.font.size=Pt(9)
def diagram(doc, rows, caption_text):
    t=doc.add_table(rows=1,cols=len(rows));t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,x in enumerate(rows):
        c=t.cell(0,i);cell_shade(c,GREEN if i%2==0 else '2F665F');cell(c,x,True,True,9.5)
    caption(doc,caption_text)

d=Document(); s=d.sections[0];s.top_margin=Inches(.68);s.bottom_margin=Inches(.66);s.left_margin=Inches(.75);s.right_margin=Inches(.75)
d.styles['Normal'].font.name='Aptos';d.styles['Normal'].font.size=Pt(10.5)
for n in ('Title','Heading 1','Heading 2','Heading 3'):
    d.styles[n].font.name='Aptos Display';d.styles[n].font.color.rgb=RGBColor(0,0,0)
d.styles['Heading 1'].font.size=Pt(17);d.styles['Heading 2'].font.size=Pt(13)

q=d.add_paragraph();q.alignment=WD_ALIGN_PARAGRAPH.CENTER;q.paragraph_format.space_before=Pt(88);r=q.add_run('AirGrid');r.bold=True;r.font.size=Pt(36);r.font.color.rgb=RGBColor(18,60,58)
q=d.add_paragraph();q.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=q.add_run('An IoT and GIS platform for verified air pollution response');r.font.size=Pt(16);r.font.color.rgb=RGBColor(35,83,76)
q=d.add_paragraph();q.alignment=WD_ALIGN_PARAGRAPH.CENTER;q.paragraph_format.space_before=Pt(30);r=q.add_run('Blue Skies Hack 2026 Detailed Project Report');r.bold=True;r.font.size=Pt(14)
q=d.add_paragraph();q.alignment=WD_ALIGN_PARAGRAPH.CENTER;q.paragraph_format.space_before=Pt(55);q.add_run('Submitted by: [Team Name]\nTeam members: [Member 1], [Member 2], [Member 3], [Member 4]\nInstitution: [Institution]\nSubmission date: 05 September 2026')
d.add_page_break()

h(d,'Executive Summary')
p(d,'AirGrid is a proposed smart-city air-pollution response system for Indian cities. It connects field IoT sensors, official/reference stations, weather and wind data, GIS mapping, hotspot analytics, and an authority action workflow. In simple terms: it helps an officer see a possible pollution problem, verify it with trusted measurements, inspect the right area, act on the likely source, and measure whether the action helped.')
p(d,'The core innovation is transparency. A physical sensor reading, a computer-model estimate, and a map interpolation are different kinds of evidence. AirGrid labels them separately and uses modelled data to guide verification—not to make unsupported enforcement claims. The current web prototype proves the map, live modelled weather and PM context, wind layer, satellite geography, source labelling, and API integration. The proposed pilot adds calibrated IoT sensors, authority data partnerships, and a verified response workflow.')
h(d,'Project Title');p(d,'AirGrid Smart IoT and GIS Platform for Air Pollution Detection Verification and Response')
h(d,'Team Introduction and Strength')
p(d,'Replace the placeholders below with final member information before submission. The proposed team structure covers the skills needed to build and validate the pilot.')
tbl(d,['Member','Proposed role','Strength'],[
['[Member 1]','Product and GIS lead','Maps user journeys, stakeholder needs and spatial decision interface'],
['[Member 2]','IoT and embedded lead','Sensor node, ESP32 firmware, connectivity, calibration support'],
['[Member 3]','Data and backend lead','Data ingestion, APIs, quality checks, hotspot analytics and database'],
['[Member 4]','Research and field lead','Environmental evidence, validation design, authorities and pilot execution'],
])
h(d,'Background of the Problem')
p(d,'Air pollution is highly local. A city average can hide a harmful pocket near a school, bus stop, construction site, traffic corridor, industrial estate, port, diesel-generator cluster, or waste-burning location. Conditions also change with wind, humidity, rainfall, traffic and emissions. The practical problem for authorities is not only “what is the AQI?” but “where should we verify first, what can we inspect, and did the response work?”')
p(d,'The National Clean Air Programme supports city-level action plans and source-apportionment studies to prioritise pollution-control actions. AirGrid is designed as the local operational layer that connects data to a documented field response.')
h(d,'Evidence and References',2)
bullets(d,[
'MoEFCC Annual Report 2024–25: NCAP action plans and source-apportionment studies. https://www.moef.gov.in/uploads/pdf-uploads/English_Annual_Report_2024-25.pdf',
'Open-Meteo Air Quality API: modelled PM and gas concentrations from CAMS data. https://open-meteo.com/en/docs/air-quality-api',
'OpenAQ API: station and measured environmental data access. https://docs.openaq.org/api',
'CPCB Air Quality Management: national real-time data and city action-plan resources. https://cpcb.nic.in/rules-7/',
])
h(d,'Problem Statement')
p(d,'What: Authorities need a reliable, local, and explainable way to turn air-quality signals into field action. Current data is frequently fragmented across apps, sensors, spreadsheets and departments. When: The issue becomes urgent during PM episodes, construction or freight peaks, industrial upsets, dry conditions, and wind-driven pollution movement. Where: The proposed pilot is Kochi–Ernakulam, with extension to Delhi NCR, Bengaluru, Mumbai and other cities.')
p(d,'Concise solution statement: AirGrid combines measured data, modelled air and weather context, IoT sensing and GIS into a source-labelled hotspot workflow that supports verification, targeted action and impact measurement.')

h(d,'Proposed Solution')
p(d,'AirGrid has five connected layers: sensing, data trust, spatial intelligence, response management, and public transparency. No single layer is treated as enough on its own. The platform is designed so a low-cost sensor can detect a change, a calibrated reference or mobile monitor can verify it, and the resulting case can be assigned to the correct authority.')
diagram(d,['Field sensing','Secure ingestion','Quality checks','GIS and hybrid grid','Case and impact tracking'],'Figure 1. End-to-end AirGrid solution block diagram')
h(d,'1 Sensing Layer',2)
tbl(d,['Data source','Examples','Purpose'],[
['Reference and official stations','CPCB CAAQMS, KSPCB, DPCC, OpenAQ','Trusted baseline and verification'],
['AirGrid IoT nodes','PM sensor, temperature/humidity/pressure sensor, GPS, ESP32','Dense local sensing around suspected hotspots'],
['Weather sources','Local weather station plus Open-Meteo','Wind direction, speed, humidity, rainfall and temperature context'],
['Mobile monitoring','Handheld/reference-grade portable monitor','Confirm a hotspot upwind and downwind'],
['Source layers','Traffic, construction, industry, waste complaints, schools and hospitals','Help inspectors choose where to investigate'],
])
h(d,'2 Data Trust Layer',2)
p(d,'Every reading carries a time, location, source, calibration state and quality flag. The platform checks freshness, physically plausible ranges, sensor uptime, missing data, and disagreement with nearby verified measurements. Suspect readings remain visible to technical users but do not drive enforcement automatically.')
h(d,'3 GIS and Hybrid Grid Layer',2)
p(d,'The city is divided into small map cells. For each cell, AirGrid starts with the modelled background air-quality estimate, then compares it with validated local stations. Nearby trusted stations correct the model more strongly than distant stations. The result is labelled as modelled, measured, corrected model, or low confidence. This makes the map useful without pretending that every point is a physical sensor.')
tbl(d,['Map label','Meaning','Appropriate use'],[
['Measured station','Direct observation at a fixed sensor','Verification and trend evidence'],
['Modelled','Atmospheric model estimate','City-wide context and gap filling'],
['Corrected model','Model adjusted by nearby verified stations','Operational prioritisation'],
['Low confidence','Sparse, stale or questionable inputs','Request verification; do not enforce from this alone'],
])
h(d,'4 Hotspot and Wind Layer',2)
p(d,'A hotspot is a priority signal, not an automatic accusation. AirGrid compares current PM2.5 with a local historical baseline, checks whether the elevation persists, looks for agreement from nearby sensors, and adds data confidence and exposure context such as schools, hospitals or dense residential areas. Wind arrows then create a downwind “risk corridor” for a field team to test with a mobile monitor. Wind identifies where to sample next; it does not prove legal source responsibility.')
diagram(d,['Current reading','Local baseline','Persistence','Nearby agreement','Wind and exposure','Priority case'],'Figure 2. Explainable hotspot scoring and response logic')
h(d,'5 Authority Action Layer',2)
p(d,'After verification, an officer creates a case. The case records the location, pollutant, confidence, probable inspection zone, owner, deadline, evidence, action taken and closure outcome. This turns a visual hotspot into an accountable task. A public view can show source-labelled status and health guidance, while internal inspection records remain restricted.')

h(d,'Hardware Architecture')
p(d,'A practical AirGrid pilot node is a rugged outdoor enclosure installed near a likely exposure location, not necessarily at a source. It measures local air, transmits data, and is periodically compared with a reference monitor. Low-cost sensors improve spatial coverage but must be calibrated; they cannot replace regulatory reference instruments.')
diagram(d,['PM sensor\nPMS7003 or SDS011','Temp humidity pressure\nBME280','Optional gas sensors\nNO2 O3 CO','ESP32 controller','LoRaWAN or 4G','Cloud API and PostGIS'],'Figure 3. AirGrid IoT node hardware block diagram')
tbl(d,['Component','Example implementation','Role in system'],[
['Microcontroller','ESP32','Reads sensors, timestamps data, buffers data during outages and sends secure payloads'],
['Particulate sensor','PMS7003, Sensirion SPS30 or equivalent calibrated PM sensor','PM1, PM2.5 and PM10 trend sensing'],
['Environmental sensor','BME280 or equivalent','Temperature, humidity and pressure for sensor correction and context'],
['Wind sensor','Cup anemometer plus wind vane, or ultrasonic weather station','Local speed and direction for plume verification'],
['Location and time','GPS module or fixed surveyed coordinates plus NTP','Reliable geotagging and clock synchronisation'],
['Connectivity','LoRaWAN gateway, NB-IoT, LTE/4G or Wi-Fi','Data transfer suited to site conditions'],
['Power','Mains with UPS; optional solar panel, charge controller and battery','Continuous operation during power interruptions'],
['Enclosure','IP65 weather-resistant box, inlet design, mounting pole','Outdoor safety, stable sampling and maintenance access'],
])
h(d,'Simplified Sensor Node Circuit',2)
p(d,'The detailed circuit is built after selecting the exact sensor model and power arrangement. A safe first prototype uses a regulated 5V supply for the PM sensor, regulated 3.3V for ESP32 and BME280, common ground, UART between ESP32 and the PM sensor, I2C for BME280, and a protected power input. For field deployment, include fuse, reverse-polarity protection, surge protection, waterproof connectors and an approved enclosure.')
tbl(d,['Block','Connection','Note'],[
['5V power input','Fuse → reverse-polarity protection → 5V rail','Use certified supply and weather-safe cabling'],
['ESP32','3.3V regulator; common ground','Runs firmware, Wi-Fi/LoRa/4G communication'],
['PM sensor','5V power; UART TX/RX to ESP32 through voltage-safe interface','Avoid directly feeding 5V UART into a 3.3V GPIO'],
['BME280','3.3V; I2C SDA/SCL to ESP32','Use pull-ups as specified by the breakout board'],
['Wind sensor','Digital pulse/voltage interface to protected GPIO/ADC','Choose interface based on anemometer model'],
['Optional purifier relay','Opto-isolated relay/contactor control','Use only with certified equipment and electrical supervision'],
])
caption(d,'Figure 4. Hardware circuit-level connection plan for a safe prototype node. Final circuitry must follow the selected component datasheets and electrical safety standards.')

h(d,'Software Technology Stack')
tbl(d,['Layer','Technology','Why it is used'],[
['Web interface','React, TypeScript, Vite, Tailwind CSS','Fast, responsive dashboard and source-labelled controls'],
['Web GIS','MapLibre GL with Leaflet fallback; satellite tiles','Interactive maps, layers, wind vectors, city search and spatial overlays'],
['Backend API','Node.js, Express, TypeScript','Secure API proxying, validation and role-based services'],
['Data ingestion','Scheduled workers, REST APIs, MQTT for IoT','Pull official/model feeds and receive sensor payloads'],
['Database','PostgreSQL plus PostGIS; TimescaleDB optional','Spatial queries, zones, map cells, time-series readings and audit history'],
['Analytics','Python or Node workers','Quality checks, calibration, baseline, hotspot and hybrid-grid jobs'],
['IoT firmware','ESP-IDF or Arduino framework for ESP32','Reliable sensor reading, buffering and telemetry'],
['Deployment','Docker, CI/CD, cloud VM/Kubernetes as scale requires','Repeatable releases, logs, backups and environment separation'],
['Security','HTTPS, secret manager, JWT/OAuth, role-based access, audit log','Protect data and authority actions'],
])
h(d,'Software Architecture Block Diagram',2)
diagram(d,['IoT nodes and APIs','MQTT REST ingestion','Raw immutable store','QA calibration engine','PostGIS and analytics','Map API cases public portal'],'Figure 5. Production software architecture')

h(d,'Optional Smart Purifier Pole')
p(d,'A sensor-equipped purifier pole may be useful at a small high-exposure micro-environment such as a bus stop, school gate, clinic entrance or indoor/semi-enclosed waiting area. It can create a local protected-air zone and offer a visible public demonstration of the system. However, it is not a city-wide pollution solution and must never be used to justify delaying source control.')
tbl(d,['Use case','How it helps','Critical limitation'],[
['Bus stop or school gate','May reduce local exposure in a defined airflow zone; sensor tracks before/after PM','Outdoor dispersion makes city-scale benefit limited'],
['Hotspot verification post','Co-located sensors show local trend and provide public alert display','Does not identify or eliminate emissions source'],
['Emergency micro-zone response','Temporary support during a severe episode','Requires certified filtration, safe maintenance and energy budget'],
])
p(d,'Recommended principle: “Control emissions first; use purification only as a local, measured exposure-reduction measure.” The pilot should compare upwind, at-pole and downwind readings before making any benefit claim. A purifier should use certified filtration, documented clean-air delivery rate, safe electrical design, filter-change schedule, noise assessment and tamper-resistant enclosure.')

h(d,'Practical Applications')
bullets(d,[
'Kochi–Ernakulam air-quality pilot: map industrial, freight, waste, construction and traffic risk zones; deploy sensors and mobile verification.',
'Construction-dust compliance: trigger a verification case when a nearby sensor shows persistent elevation and wind points from an active work site.',
'School and hospital protection: warn administrators when verified PM rises, identify safer outdoor activity windows, and prioritise micro-zone interventions.',
'Smart-city control room: merge official stations, local IoT network, complaints and field actions into one auditable workspace.',
'Port and industrial-estate monitoring: locate monitoring posts along prevailing wind directions and verify downwind impact patterns.',
])

h(d,'Methodology and Workflow')
tbl(d,['Stage','Implementation method','Output'],[
['Collect','Read API data and IoT payloads with source metadata','Raw measured/modelled records'],
['Validate','Freshness, calibration, outlier, uptime and neighbour checks','Quality flags and trusted dataset'],
['Map','Build source-labelled layers and corrected grid','GIS view with confidence'],
['Detect','Compare current values with baseline and persistence','Ranked hotspot signal'],
['Verify','Mobile monitor samples upwind/downwind; officer documents evidence','Confirmed or rejected case'],
['Respond','Assign action, owner, deadline and compliance checks','Authority action record'],
['Measure','Compare before/after with weather and control location','Impact and repeat-event metrics'],
])
h(d,'Connection to Sustainability')
tbl(d,['SDG','AirGrid contribution'],[
['SDG 3 Good Health and Well Being','Faster verification and health-risk communication for vulnerable groups'],
['SDG 9 Industry Innovation and Infrastructure','Local digital environmental infrastructure built around interoperable data'],
['SDG 11 Sustainable Cities and Communities','Operational tools for cleaner transport, waste, construction and neighbourhood management'],
['SDG 13 Climate Action','Weather-aware planning and evidence for low-emission interventions'],
])
p(d,'The expected early impact is better decision quality: fewer blind inspections, quicker verification, clearer public communication and documented intervention outcomes. Health and economic benefit claims will be measured only after the pilot collects validated data.')

h(d,'Implementation Timeline')
tbl(d,['Phase','Duration','Key work and output'],[
['1 Build and bench test','Weeks 1–4','Assemble 3–5 nodes; validate sensors beside a trusted reference monitor; complete enclosure and firmware'],
['2 Kochi pilot setup','Months 2–3','Partner agreements, site survey, gateway/connectivity, source inventory and baseline data'],
['3 Field pilot','Months 4–6','Install 10–20 calibrated nodes, weather station and mobile verification protocol'],
['4 Analytics validation','Months 6–8','Calibrate grid, test hotspot precision, measure false positives and refine workflow'],
['5 Authority integration','Months 8–10','Case management, roles, inspection evidence, public view and reporting'],
['6 Scale','Months 10–18','Extend to Delhi NCR, Mumbai, Bengaluru and additional sensor/source layers'],
])
h(d,'Requirements and Collaboration Request')
tbl(d,['Need','Requested support from Amrita TBI TEC and partners'],[
['Mentorship','Environmental sensing, calibration, IoT productisation, GIS, procurement and government collaboration'],
['Pilot partnership','Kochi Corporation, Kerala PCB, academic laboratories, school/hospital and community partners'],
['Hardware and field resources','Reference-monitor access, 10–20 sensor nodes, weather station, LoRa/4G connectivity and mobile monitor'],
['Technical infrastructure','Cloud credits, PostGIS database, secure secrets, monitoring and deployment support'],
['Funding','Prototype components, rugged enclosures, calibration, field travel, connectivity and maintenance'],
['Validation support','Independent testing protocol and expert review before any authority action or public claim'],
])
p(d,'Readiness: The team is prepared to collaborate with Amrita TBI or TEC on a scoped Kochi pilot. The first deliverable would be a jointly approved sensor-calibration and verification protocol, followed by an evidence-based deployment plan.')

h(d,'Prototype Proof of Concept')
p(d,'The current AirGrid prototype demonstrates the software foundation: source-labelled live Open-Meteo PM and weather context; India-wide GIS map with wind vectors and satellite geography; optional OpenAQ station discovery; Kochi/Delhi location navigation; API refresh and failure states; and a clear Live/Demo switch. The demonstration mode is used only to illustrate the future workflow, while the real proposed system is the calibrated IoT and verified-station architecture described in this report.')
tbl(d,['Recommended report screenshot','Short caption'],[
['Map Workspace in Live mode','Nationwide modelled PM surface, wind direction and satellite geography, all source-labelled'],
['Live Data page','Open-Meteo weather/air context and separate OpenAQ station evidence'],
['Kochi city view','Example pilot geography for field-sensor deployment and response workflow'],
['Live/Demo switch','Transparency feature: separates implementation proof from simulated training data'],
])

h(d,'Innovator Perspective')
h(d,'Zoom In',2)
p(d,'At one busy Kochi location, a sensor detects that PM2.5 is rising. AirGrid checks whether the sensor is healthy, compares it with a reference station and modelled conditions, shows the wind direction, and guides the officer to measure upwind and downwind. The team can then inspect likely local sources and record action.')
h(d,'Zoom Out',2)
p(d,'Across a city, many small verified observations become a living pollution intelligence network. Instead of relying only on a city-wide average or a black-box prediction, AirGrid helps cities build a transparent evidence trail from sensing to enforcement to measurable improvement. The long-term vision is scalable clean-air infrastructure for Indian cities, with local data ownership and public trust.')
h(d,'Conclusion')
p(d,'AirGrid is a feasible, scalable and socially useful air-pollution response platform. Its immediate value is not that it promises perfect prediction; it provides a disciplined method for using data responsibly. With calibrated sensors, trusted partners, GIS and a measurable authority workflow, the platform can help Kochi move from identifying pollution episodes to verifying and reducing local exposure.')
h(d,'References')
for x in [
'Ministry of Environment Forest and Climate Change. Annual Report 2024–25. https://www.moef.gov.in/uploads/pdf-uploads/English_Annual_Report_2024-25.pdf',
'Open-Meteo. Air Quality API Documentation. https://open-meteo.com/en/docs/air-quality-api',
'OpenAQ. API Documentation. https://docs.openaq.org/api',
'Central Pollution Control Board. Air Quality Management. https://cpcb.nic.in/rules-7/',
'Ministry of Environment Forest and Climate Change. Annual Report 2023–24 and NCAP overview. https://moef.gov.in/uploads/2023/05/Annual-Report-English-2023-24.pdf',
]:p(d,x)
d.save(OUT);print(OUT)
